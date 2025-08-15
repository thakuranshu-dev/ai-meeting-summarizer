from docx import Document
import os
import tempfile

def export_to_docx(summary, action_items):
    """
    Export meeting summary and action items to a DOCX file.
    Returns file path.
    """
    doc = Document()
    doc.add_heading("Meeting Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Action Items", level=1)
    for item in action_items:
        doc.add_paragraph(f"- {item}")

    tmp_dir = tempfile.gettempdir()
    file_path = os.path.join(tmp_dir, "meeting_report.docx")
    doc.save(file_path)
    return file_path
